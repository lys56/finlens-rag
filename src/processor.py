"""
据解析与清洗模块
功能：
1. 遍历 my_rag_data 目录下的所有文件
2. 解析 PDF (使用 pdfplumber 提取文本，保留页码信息)
3. 解析 Word (使用 python-docx 提取文本，保留段落信息)
4. 解析 HTML (使用 BeautifulSoup 去噪)
5. 解析 TXT (纯文本文件)
6. 将清洗后的文本保存到 data/processed 目录，并保留页码元数据
"""

import os
import re
import json
import pdfplumber
from bs4 import BeautifulSoup
from pathlib import Path
from tqdm import tqdm

# 尝试导入 Word 处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word 文档解析功能将不可用")

class FinancialDataProcessor:
    def __init__(self, input_dir="my_rag_data", output_dir="data/processed"):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        # 用于跟踪已处理的文件（基于文件修改时间）
        self.processed_cache = {}

    def clean_text(self, text: str) -> str:
        """通用文本清洗函数"""
        if not text:
            return ""
            
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'^\s*[-_]?\s*\d+\s*[-_]?\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def _should_process(self, file_path: Path) -> bool:
        """检查文件是否需要处理（基于修改时间）"""
        output_file = self.output_dir / f"{file_path.stem}.json"
        if not output_file.exists():
            return True
        # 如果源文件比输出文件新，需要重新处理
        return file_path.stat().st_mtime > output_file.stat().st_mtime

    def parse_pdf_with_pages(self, file_path: Path) -> list:
        """
        解析 PDF 文件，返回带页码的文本列表
        返回格式: [{"page": 1, "text": "..."}, {"page": 2, "text": "..."}, ...]
        """
        pages_data = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        cleaned_text = self.clean_text(text)
                        if cleaned_text:
                            pages_data.append({
                                "page": page_num,
                                "text": cleaned_text
                            })
            return pages_data
        except Exception as e:
            print(f"警告: PDF 解析失败 {file_path.name}: {e}")
            return []

    def parse_docx_with_paragraphs(self, file_path: Path) -> list:
        """
        解析 Word (.docx) 文件，返回带段落信息的文本列表
        返回格式: [{"page": 1, "text": "..."}, {"page": 2, "text": "..."}, ...]
        注意：Word 文档没有明确的"页"概念，这里用段落分组来模拟分页
        """
        if not DOCX_AVAILABLE:
            print("警告: Word 解析不可用: python-docx 未安装")
            return []
        
        paragraphs_data = []
        try:
            doc = Document(file_path)
            
            # 将文档按段落分组，每 N 个段落为一"页"（模拟分页）
            # 或者按段落长度累计，达到一定字符数后分页
            current_page = 1
            current_text = []
            current_length = 0
            chars_per_page = 2000  # 每页约 2000 字符
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if not para_text:
                    continue
                
                # 如果当前页已满，创建新页
                if current_length + len(para_text) > chars_per_page and current_text:
                    paragraphs_data.append({
                        "page": current_page,
                        "text": "\n".join(current_text)
                    })
                    current_page += 1
                    current_text = []
                    current_length = 0
                
                current_text.append(para_text)
                current_length += len(para_text)
            
            # 添加最后一页
            if current_text:
                paragraphs_data.append({
                    "page": current_page,
                    "text": "\n".join(current_text)
                })
            
            # 处理表格
            for table in doc.tables:
                table_texts = []
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells]
                    table_texts.append(" | ".join(row_texts))
                
                if table_texts:
                    table_content = "\n".join(table_texts)
                    # 如果表格较大，单独作为一页
                    if len(table_content) > 500:
                        paragraphs_data.append({
                            "page": current_page,
                            "text": f"[表格]\n{table_content}"
                        })
                        current_page += 1
                    else:
                        # 小表格合并到当前页
                        if paragraphs_data:
                            paragraphs_data[-1]["text"] += f"\n\n[表格]\n{table_content}"
            
            # 清洗所有页面的文本
            for item in paragraphs_data:
                item["text"] = self.clean_text(item["text"])
            
            return paragraphs_data
            
        except Exception as e:
            print(f"警告: Word 解析失败 {file_path.name}: {e}")
            return []

    def parse_html(self, file_path: Path) -> str:
        """解析 HTML 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
                
                for script in soup(["script", "style", "head", "title", "meta"]):
                    script.extract()
                
                text = soup.get_text(separator='\n')
                return text
        except Exception as e:
            print(f"警告: HTML 解析失败 {file_path.name}: {e}")
            return ""
    
    def parse_txt(self, file_path: Path) -> str:
        """解析纯文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            print(f"警告: TXT 解析失败 {file_path.name}: {e}")
            return ""

    def process_all(self, only_new_files=False):
        """
        主处理流程 - 支持多种文件类型
        Args:
            only_new_files: 如果为 True，只处理新文件或已修改的文件
        """
        print("开始处理数据...")
        print(f"输入目录: {self.input_dir}")
        print(f"输出目录: {self.output_dir}")
        
        all_files = list(self.input_dir.rglob("*"))
        files_to_process = [f for f in all_files if f.is_file()]
        
        # 如果只处理新文件，过滤已处理的文件
        if only_new_files:
            files_to_process = [
                f for f in files_to_process 
                if self._should_process(f)
            ]
            if not files_to_process:
                print("没有需要处理的新文件")
                return
        
        # 支持的文件类型
        supported_extensions = {
            '.pdf': 'PDF',
            '.docx': 'Word',
            '.doc': 'Word (旧格式)',
            '.html': 'HTML',
            '.htm': 'HTML',
            '.txt': 'Text'
        }
        
        success_count = 0
        skipped_count = 0
        error_count = 0
        
        for file_path in tqdm(files_to_process, desc="Processing"):
            file_ext = file_path.suffix.lower()
            
            # 跳过不支持的文件类型
            if file_ext not in supported_extensions:
                skipped_count += 1
                continue
            
            try:
                if file_ext == '.pdf':
                    # PDF 文件：保存为带页码的 JSON 格式
                    pages_data = self.parse_pdf_with_pages(file_path)
                    if pages_data:
                        output_filename = f"{file_path.stem}.json"
                        output_path = self.output_dir / output_filename
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            json.dump({
                                "source": file_path.name,
                                "file_type": "pdf",
                                "pages": pages_data
                            }, f, ensure_ascii=False, indent=2)
                        success_count += 1
                    else:
                        error_count += 1
                        
                elif file_ext == '.docx':
                    # Word 文件：保存为带段落信息的 JSON 格式
                    paragraphs_data = self.parse_docx_with_paragraphs(file_path)
                    if paragraphs_data:
                        output_filename = f"{file_path.stem}.json"
                        output_path = self.output_dir / output_filename
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            json.dump({
                                "source": file_path.name,
                                "file_type": "docx",
                                "pages": paragraphs_data  # 使用相同的 pages 字段保持兼容
                            }, f, ensure_ascii=False, indent=2)
                        success_count += 1
                    else:
                        error_count += 1
                
                elif file_ext == '.doc':
                    # 旧版 Word 格式 (.doc) - 需要额外库支持，暂时跳过或提示
                    print(f"警告: 暂不支持旧版 Word 格式 (.doc): {file_path.name}")
                    print("   提示: 请将文件转换为 .docx 格式")
                    skipped_count += 1
                    
                elif file_ext in ['.html', '.htm']:
                    # HTML 文件：保存为普通文本
                    content = self.parse_html(file_path)
                    cleaned_content = self.clean_text(content)
                    
                    if cleaned_content:
                        # 将 HTML 转换为带页码的格式（按段落分组）
                        paragraphs = cleaned_content.split('\n\n')
                        pages_data = []
                        current_page = 1
                        current_text = []
                        chars_per_page = 2000
                        current_length = 0
                        
                        for para in paragraphs:
                            if current_length + len(para) > chars_per_page and current_text:
                                pages_data.append({
                                    "page": current_page,
                                    "text": "\n\n".join(current_text)
                                })
                                current_page += 1
                                current_text = []
                                current_length = 0
                            
                            current_text.append(para)
                            current_length += len(para)
                        
                        if current_text:
                            pages_data.append({
                                "page": current_page,
                                "text": "\n\n".join(current_text)
                            })
                        
                        output_filename = f"{file_path.stem}.json"
                        output_path = self.output_dir / output_filename
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            json.dump({
                                "source": file_path.name,
                                "file_type": "html",
                                "pages": pages_data
                            }, f, ensure_ascii=False, indent=2)
                        success_count += 1
                    else:
                        error_count += 1
                        
                elif file_ext == '.txt':
                    # TXT 文件：保存为带页码的格式
                    content = self.parse_txt(file_path)
                    cleaned_content = self.clean_text(content)
                    
                    if cleaned_content:
                        # 按段落分组
                        paragraphs = cleaned_content.split('\n\n')
                        pages_data = []
                        current_page = 1
                        current_text = []
                        chars_per_page = 2000
                        current_length = 0
                        
                        for para in paragraphs:
                            if current_length + len(para) > chars_per_page and current_text:
                                pages_data.append({
                                    "page": current_page,
                                    "text": "\n\n".join(current_text)
                                })
                                current_page += 1
                                current_text = []
                                current_length = 0
                            
                            current_text.append(para)
                            current_length += len(para)
                        
                        if current_text:
                            pages_data.append({
                                "page": current_page,
                                "text": "\n\n".join(current_text)
                            })
                        
                        output_filename = f"{file_path.stem}.json"
                        output_path = self.output_dir / output_filename
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            json.dump({
                                "source": file_path.name,
                                "file_type": "txt",
                                "pages": pages_data
                            }, f, ensure_ascii=False, indent=2)
                        success_count += 1
                    else:
                        error_count += 1
                        
            except Exception as e:
                print(f"处理文件失败 {file_path.name}: {e}")
                error_count += 1
        
        print("-" * 50)
        print("处理完成!")
        print(f"   ✓ 成功: {success_count} 个文件")
        print(f"   ⊘ 跳过: {skipped_count} 个文件（不支持的类型）")
        print(f"   ✗ 失败: {error_count} 个文件")
        print(f"查看结果: {self.output_dir}")

if __name__ == "__main__":
    processor = FinancialDataProcessor(
        input_dir="my_rag_data",
        output_dir="data/processed"
    )
    processor.process_all()
